"""
DOCX Parser — Extract text from .docx Word documents.
Uses python-docx to extract paragraphs, headings, tables, and metadata.
"""

from __future__ import annotations

import logging
from pathlib import Path
from typing import Any

logger = logging.getLogger(__name__)

SUPPORTED_EXTENSIONS = {".docx"}


def can_parse(file_type: str) -> bool:
    return file_type.lower().lstrip(".") in {ext.lstrip(".") for ext in SUPPORTED_EXTENSIONS}


async def parse(file_path: str) -> dict[str, Any]:
    """
    Parse a DOCX file.

    Returns:
        dict with keys: text, metadata, pages, sections
    """
    from docx import Document

    path = Path(file_path)
    logger.info(f"Parsing DOCX: {path.name}")

    try:
        doc = Document(file_path)

        # Extract core properties
        props = doc.core_properties
        doc_metadata = {
            "format": "docx",
            "title": props.title or "",
            "author": props.author or "",
            "subject": props.subject or "",
            "created": str(props.created) if props.created else "",
            "modified": str(props.modified) if props.modified else "",
        }

        # Extract text by paragraphs, tracking sections
        sections = []
        current_section = {"title": "", "level": 0, "text": ""}
        all_text_parts = []

        for para in doc.paragraphs:
            text = para.text.strip()
            if not text:
                continue

            # Check if it's a heading
            style_name = (para.style.name or "").lower()
            if style_name.startswith("heading"):
                # Save current section
                if current_section["text"].strip():
                    sections.append(current_section)

                # Determine heading level
                try:
                    level = int(style_name.replace("heading", "").strip())
                except ValueError:
                    level = 1

                current_section = {"title": text, "level": level, "text": ""}
            else:
                current_section["text"] += text + "\n"

            all_text_parts.append(text)

        # Don't forget the last section
        if current_section["text"].strip() or current_section["title"]:
            sections.append(current_section)

        # Extract tables as text
        table_texts = []
        for table in doc.tables:
            rows_text = []
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells]
                rows_text.append(" | ".join(cells))
            table_text = "\n".join(rows_text)
            table_texts.append(table_text)
            all_text_parts.append(table_text)

        full_text = "\n\n".join(all_text_parts)

        doc_metadata["paragraph_count"] = len(doc.paragraphs)
        doc_metadata["table_count"] = len(doc.tables)
        doc_metadata["section_count"] = len(sections)
        doc_metadata["char_count"] = len(full_text)
        doc_metadata["headings"] = [s["title"] for s in sections if s["title"]]

        return {
            "text": full_text,
            "metadata": doc_metadata,
            "pages": [{"page": 1, "text": full_text}],  # DOCX doesn't have page breaks
            "sections": sections,
        }

    except ImportError:
        raise ValueError("python-docx not installed. Run: pip install python-docx")
    except Exception as e:
        logger.error(f"DOCX parsing failed for {path.name}: {e}")
        raise ValueError(f"Failed to parse DOCX: {e}") from e
